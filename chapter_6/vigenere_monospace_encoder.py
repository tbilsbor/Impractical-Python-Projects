#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Jan 15 15:12:08 2019

@author: toddbilsborough

Chapter 6 challenge project: Vigenere cipher using monospace font
from Impractical Python Projects

Objective
- As per the vigenere_cipher_encoder, but hiding the message in a monospace
font between words

"""

import vigenere_cipher_encoder as vce
import docx
from docx.shared import RGBColor
import sys

def verify_spaces(fake_text, real_text):
    """Verifies that there are enough spaces in the fake text
    to hold the message.
    Takes a string rather than a list for the real_text"""
    space_count = 0
    for line in fake_text:
        space_count += line.count(" ")
    char_count = len(real_text)
    if space_count < char_count:
        print("\nInsufficient spaces in fake message")
        print("Terminating")
        sys.exit()
        
def format_text(paragraph, white=False):
    """Puts the text in the correct font and color"""
    run = paragraph.runs[-1]
    font = run.font
    font.name = "Courier New"
    if white:
        font.color.rgb = RGBColor(255, 255, 255)

def main():
    """Load the fake and real messages as lists.
    Turn the real message into a cipher.
    Add letterhead to new doc.
    Interleave the fake message and the cipher."""
    fake_list = vce.load_fake_text('fakeMessage.docx')
    #real_list = vce.load_real_text('realMessage.docx')
    real_list = ["This is a hidden message"]
    key = "splendiferous"
    cipher_list = vce.encode_vigenere(real_list, key)
    # Turn cipher_list into a single string
    ciphertext = ""
    for line in cipher_list:
        ciphertext += "{} ".format(line)
    verify_spaces(fake_list, ciphertext)
    # Reverse and listify the cipher text for easy popping
    ciphertext = list(ciphertext[::-1])
    
    # Add letterhead to template
    doc = docx.Document('template_monospace.docx')
    doc.add_heading('Morland Holmes', 0) # 0 for max heading
    # 1 for subtitle heading
    subtitle = doc.add_heading('Global Consulting & Negotiations', 1)
    subtitle.alignment = 1
    doc.add_heading('', 1) 
    paragraph = doc.add_paragraph('January 15, 2019')
    format_text(paragraph)
    doc.add_paragraph('')
    
    # interleave fake message words with cipher list characters
    for line in fake_list:
        paragraph = doc.add_paragraph('')
        if line == "":
            continue
        running_text = "" # Builds up a run of characters from the fake
        for char in line:
            if char == " ":
                # Add the fake text that has accumulated so far
                paragraph.add_run(running_text)
                format_text(paragraph)
                running_text = ""
                # Make a new run for the character from the message
                if len(ciphertext) > 0: # While there's still message left
                    if ciphertext[-1] == " ": # Maintain spaces from message
                        ciphertext.pop()
                        running_text += " "
                    else:
                        paragraph.add_run(ciphertext.pop())
                        format_text(paragraph, white=True)
                else: # If you're out of message, spaces stay spaces
                    running_text += " "
            else:
                running_text += char
        # Add the last word from the fake text
        paragraph.add_run(running_text)
        format_text(paragraph)
        running_text = ""
        # Left justify or you get weird spacing issues
        paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT
        vce.set_spacing(paragraph)
    doc.save('ciphertext_message_monospace.docx')
    print("Done")
    
if __name__ == '__main__':
    main()
